# hack nsa por favor
# need money for things
# send money please: paypal => fureurfatal@gmail.com

from docx import Document
from time import strftime
import schedule
import argparse
import locale
import docx
import time
import json
import bz2
import os

parser = argparse.ArgumentParser(description='Diary app.')
parser.add_argument('--wd', action='store', type=str, required=True)
args = parser.parse_args()


# SAVING FILE LOCATION:
CONFIG_FILE = args.wd + '/config.json'
JSON_FILE = args.wd+'/src/data.json'
WORD_FILE_LOCATION = args.wd+'/src/temp_file.docx'

locale.setlocale(locale.LC_ALL, "de_DE")
debug = False

with open(CONFIG_FILE, 'r') as file:
    j = ''.join(file.readlines())
    config = json.loads(j)


def log(text, l):
    l.write('(' + strftime('%d/%m/%Y') + ') ')
    l.write('['+strftime('%H:%M:%S')+'] ')
    l.write(text+"\n")
    # log('init successfull')
def create_template(l=False):
    if l != False:
        log('created the document', l)
    doc = Document()
    date = strftime('%d %B %Y')

    # title
    doc.add_heading(f"Tagebuch ({date})", 0)

    p1 = doc.add_paragraph('')
    e = p1.add_run('Was ich heute gemacht habe\n')
    e.bold = True
    e.font.name = 'Calibri'

    p2 = doc.add_paragraph('')
    e = p2.add_run('Was ich grade mache\n')
    e.bold = True
    e.font.name = 'Calibri'

    doc.save(WORD_FILE_LOCATION)
    log('document was sucessfully created', l)

def read_file():
    doc = docx.Document(WORD_FILE_LOCATION)
    ft = []
    for p in doc.paragraphs:
        ft.append(p.text)
    date = ft[0].replace('Tagebuch (', '').replace(')', '')
    q1 = []
    for i in range(len(ft)-1):
        if (ft[i+1].replace('Was ich grade mache\n', '') != ft[i+1]):
            break
        else:
            q1.append(ft[i+1].replace('Was ich heute gemacht habe\n',''))
    q1 = '\n'.join(q1)
    q2 = [j.replace('Was ich grade mache\n', '') for j in ft[i+1:]]
    q2 = '\n'.join(q2)
    return (date, q1, q2)

def save_question():
    (date, q1, q2) = read_file()
    with open(JSON_FILE, 'r') as file:
        lines = file.readlines()
        if lines != []:
            Jfile = json.loads(''.join(lines))
        else:
            Jfile = []
    Jfile.append({
        'Date': date,
        'Q1': q1,
        'Q2': q2
    })
    with open(JSON_FILE, 'w') as file:
        json.dump(Jfile, file, indent=2)
    with open(CONFIG_FILE, 'r') as file:
        j = ''.join(file.readlines())
        conf = json.loads(j)
    conf['written'] = True
    with open(CONFIG_FILE, 'w') as file:
        json.dump(conf, file, indent=2)

def getTasks(name):
    # funcion from https://www.bogotobogo.com/python/python-Windows-Check-if-a-Process-is-Running-Hanging-Schtasks-Run-Stop.php
    r = os.popen('tasklist /v').read().strip().split('\n')
    for i in range(len(r)):
        s = r[i]
        if name in r[i]:
            return r[i]
    return []

def is_word_open():
    r = getTasks('WINWORD.EXE')
    if not r:
        return False
    else:
        return True

# App
def mainn():
    logger = open(args.wd+'/src/log.txt', 'a')
    log('successfully opened the file (probably startup)', logger)
    create_template(logger)
    os.startfile(WORD_FILE_LOCATION)
    while(is_word_open()):
        time.sleep(5)
        log('sleepy sleepy for... ... ...5 ...seconds', logger)
    log('saving the questions...', logger)
    save_question()
    log('saved the file, waiting for next startup')

def check():
    with open(CONFIG_FILE, 'r') as file:
        j = ''.join(file.readlines())
        conf = json.loads(j)
    if conf['written'] == False:
        mainn()
    else:
        conf['written'] = False
        with open(CONFIG_FILE, 'w') as file:
            json.dump(conf, file, indent=2)
# mainn()
if __name__ == '__main__':
    # mainn()
    logger = open(args.wd+'/src/log.txt', 'a')
    log('just opened the file', logger)
    schedule.every().day.at(config['time']).do(mainn)
    schedule.every().day.at(config['checkupTime']).do(check)
    log(f"scheduled opener for {config['time']} and checkup for {config['checkupTime']}", logger)
    logger.close()

    while True:
        schedule.run_pending()
        time.sleep(40)
