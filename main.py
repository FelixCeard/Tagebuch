# hack nsa por favor
# need money for things
# send money please: paypal => fureurfatal@gmail.com

from docx import Document
from docx.shared import *
from time import strftime
import time
import docx
import locale
import json
import bz2
import os
import schedule

# SAVING FILE LOCATION:
JSON_FILE = 'data.json'
WORD_FILE_LOCATION = 'temp_file.docx'
CONFIG_FILE = 'config.json'

locale.setlocale(locale.LC_ALL, "de_DE")

def create_template():
    doc = Document()
    date = strftime('%d %B %Y')

    # title
    doc.add_heading(f"Tagebuch ({date})", 0)

    p1 = doc.add_paragraph('')
    e = p1.add_run('Was ich heute gemacht habe\n')
    e.bold = True
    e.font.name = 'Calibri'

    p2 = doc.add_paragraph('')
    e = p2.add_run('Was ich grade mache\n')
    e.bold = True
    e.font.name = 'Calibri'

    doc.save(WORD_FILE_LOCATION)

def read_file():
    doc = docx.Document(WORD_FILE_LOCATION)
    ft = []
    for p in doc.paragraphs:
        ft.append(p.text)
    date = ft[0].replace('Tagebuch (', '').replace(')', '')
    q1 = []
    for i in range(len(ft)-1):
        if (ft[i+1].replace('Was ich grade mache\n', '') != ft[i+1]):
            break
        else:
            q1.append(ft[i+1].replace('Was ich heute gemacht habe\n',''))
    q1 = '\n'.join(q1)
    q2 = [j.replace('Was ich grade mache\n', '') for j in ft[i+1:]]
    q2 = '\n'.join(q2)
    return (date, q1, q2)

def save_question():
    (date, q1, q2) = read_file()
    with open(JSON_FILE, 'r') as file:
        lines = file.readlines()
        if lines != []:
            Jfile = json.loads(''.join(lines))
        else:
            Jfile = []
    Jfile.append({
        'Date': date,
        'Q1': q1,
        'Q2': q2
    })
    with open(JSON_FILE, 'w') as file:
        json.dump(Jfile, file)
    with open(CONFIG_FILE, 'r') as file:
        j = ''.join(file.readlines())
        conf = json.loads(j)
    conf['written'] = True
    with open(CONFIG_FILE, 'w') as file:
        json.dump(conf, file)

def getTasks(name):
    # funcion from https://www.bogotobogo.com/python/python-Windows-Check-if-a-Process-is-Running-Hanging-Schtasks-Run-Stop.php
    r = os.popen('tasklist /v').read().strip().split('\n')
    for i in range(len(r)):
        s = r[i]
        if name in r[i]:
            return r[i]
    return []

def is_word_open():
    r = getTasks('WINWORD.EXE')
    if not r:
        return False
    else:
        return True

# App
def main():
    create_template()
    os.startfile(WORD_FILE_LOCATION)
    while(is_word_open()):
        time.sleep(5)
        print('debug oh yeah')
    print('saving the questions...')
    save_question()

# conf
with open(CONFIG_FILE, 'r') as file:
    j = ''.join(file.readlines())
    conf = json.loads(j)
schedule.every().day.at(conf['time']).do(main)
print(f"scheduled for {conf['time']}")

while True:
        schedule.run_pending()
        time.sleep(20)
